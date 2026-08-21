"""data/build_corpus.py — 从笔记文档构建训练语料

把 `笔记/贵系.docx`、`笔记/量化.docx`、`笔记/宇宙与物理规律.docx`
三个文档的正文抽取合并，生成 `training_data.txt`（字符级 GPT 的训练语料）。

用法
----
.. code-block:: bash

    python -m data.build_corpus            # 相对项目根
    python data/build_corpus.py            # 直接运行
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

HERE = Path(__file__).resolve().parent
PROJECT = HERE.parent  # 本项目根
REPO = PROJECT.parent.parent  # 仓库根
DOCS = [
    REPO / "笔记" / "贵系.docx",
    REPO / "笔记" / "量化.docx",
    REPO / "笔记" / "宇宙与物理规律.docx",
]
OUT = PROJECT / "training_data.txt"


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # 也抽取表格文本
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras)


def main() -> None:
    parts = []
    for d in DOCS:
        if not d.exists():
            print(f"⚠ 缺少文档: {d}")
            continue
        text = extract_docx(d)
        parts.append(text)
        print(f"✓ {d.name}: {len(text)} 字符")
    corpus = "\n\n".join(parts)
    OUT.write_text(corpus, encoding="utf-8")
    print(f"语料已写入 {OUT}：{len(corpus)} 字符")


if __name__ == "__main__":
    main()
